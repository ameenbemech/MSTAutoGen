import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor

# ---------------- STREAMLIT UI ----------------
st.set_page_config(page_title="MST AutoGen (Amana AQAI)", layout="wide")

st.markdown("""
<div style="background-color:#1E90FF;padding:12px;border-radius:10px;text-align:center;">
<h2 style="color:white;font-style:italic;">📑 METHOD STATEMENT AUTOGENERATION (AMANA AQAI)</h2>
</div>
""", unsafe_allow_html=True)
st.markdown("---")

col1, col2 = st.columns(2)
with col1:
    project_name = st.text_input("🏗️ Project Name")
with col2:
    activity_title = st.text_input("⚙️ Activity Title")

client_name = st.text_input("👤 Client Name")
logo_file = st.file_uploader("🖼️ Upload Client Logo", type=["png", "jpg", "jpeg"])
template_file = st.file_uploader("📄 Upload Word Template (.docx)", type=["docx"])
uploaded_specs = st.file_uploader("Upload Project Specs (PDFs)", type=["pdf"], accept_multiple_files=True)

# ---------------- OPENAI API KEY ----------------
OPENAI_API_KEY = "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"  # <-- Replace with your key
client = OpenAI(api_key=OPENAI_API_KEY)

# ---------------- RAG VECTOR ----------------
@st.cache_resource
def build_vectorstore(files):
    docs = []
    for file in files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.name)[1]) as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name

        if file.name.endswith(".pdf"):
            loader = PyPDFLoader(tmp_path)
        else:
            loader = Docx2txtLoader(tmp_path)
        docs.extend(loader.load())

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = text_splitter.split_documents(docs)
    embeddings = HuggingFaceEmbeddings()
    return FAISS.from_documents(split_docs, embeddings)

rag_index = build_vectorstore(uploaded_specs) if uploaded_specs else None

# ---------------- AI + RAG CONTENT ----------------
def generate_rag_ai_content(activity_title, section_title, rag_index):
    """Generate 10 bullet points combining AI knowledge and RAG content"""
    rag_snippets = ""
    if rag_index:
        results = rag_index.similarity_search(f"{activity_title} {section_title}", k=3)
        rag_snippets = "\n".join([r.page_content for r in results])

    prompt = f"""
You are an expert in Testing & Commissioning and Method Statements.
Generate 10 concise bullet points for "{activity_title}" covering "{section_title}".
Use both your own knowledge and the factual details provided below.
---
{rag_snippets}
---
The bullet points should be practical, clear, and written in simple English.
Do not include any leading phrases like "Certainly!" or "Here are...".
"""

    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    output = completion.choices[0].message.content
    return [line.strip("•- ").strip() for line in output.split("\n") if line.strip()]

# ---------------- INSERT LOGO ----------------
def insert_logo_header(doc, logo_file, client_name):
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()

    if logo_file:
        # ✅ Save logo properly without deleting before doc.save()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_logo:
            tmp_logo.write(logo_file.getbuffer())
            tmp_logo.flush()
            logo_path = tmp_logo.name

        # Insert logo
        run.add_picture(logo_path, width=Inches(1.2))

        # ✅ Delete AFTER saving the Word document (handle at the end)
        doc._logo_temp_path = logo_path  # store path in doc object

    if client_name:
        run.add_text(f"  {client_name}")

# ---------------- SIMPLE SECTION TITLES ----------------
section_ph = {
    "{{Standards & References}}": "Standards & References",
    "{{Tools & Equipments}}": "Tools & Equipments",
    "{{Installation Steps}}": "Installation Steps",
    "{{Pre-Commissioning}}": "Pre-Commissioning",
    "{{Commissioning Steps}}": "Commissioning Steps",
    "{{QA/QC Records}}": "QA/QC Records",
    "{{Safety}}": "Safety"
}

# ---------------- GENERATE METHOD STATEMENT ----------------
if st.button("🚀 Generate Method Statement"):
    if not template_file:
        st.error("Please upload a Word template.")
        st.stop()

    doc = Document(template_file)
    insert_logo_header(doc, logo_file, client_name)

    # Replace Activity and Project placeholders
    for para in doc.paragraphs:
        para.text = para.text.replace("{{Activity_Title}}", activity_title)
        para.text = para.text.replace("{{Project_Name}}", project_name)

    # Process section placeholders
    placeholders_to_process = []
    for para in doc.paragraphs:
        for ph, section_title in section_ph.items():
            if ph in para.text:
                placeholders_to_process.append((para, ph, section_title))

    progress = st.progress(0)
    total_sections = len(placeholders_to_process)
    for idx, (para, ph, section_title) in enumerate(placeholders_to_process):
        para.text = para.text.replace(ph, "")
        bullets = generate_rag_ai_content(activity_title, section_title, rag_index)
        # Alternate colors: AI = blue, RAG = green
        for i, b in enumerate(bullets):
            run = para.add_run(f"• {b}\n")
            run.font.size = Pt(11)
            # First 5 bullets AI (blue), next 5 bullets RAG (green)
            if i < 5:
                run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                run.font.color.rgb = RGBColor(0, 128, 0)
        progress.progress((idx + 1) / total_sections)

    import time

    # ---------------- SAFE SAVE AND VALIDATION ----------------
    output_path = os.path.join(
        tempfile.gettempdir(),
        f"Method_Statement_{activity_title.replace(' ', '_')}.docx"
    )

    # Save DOCX
    doc.save(output_path)

    # ✅ Clean up temp logo only after saving
    if hasattr(doc, "_logo_temp_path") and os.path.exists(doc._logo_temp_path):
        time.sleep(0.2)  # tiny delay for cloud sync
        os.remove(doc._logo_temp_path)


    # ✅ Validate that Word file is readable before offering download
    def validate_docx(path):
        try:
            test_doc = Document(path)
            _ = [p.text for p in test_doc.paragraphs]  # try reading paragraphs
            return True
        except Exception as e:
            st.error(f"⚠️ Generated Word file is not readable: {e}")
            return False


    if validate_docx(output_path):
        st.success(f"✅ Method statement for '{activity_title}' generated successfully!")

        # Download button
        with open(output_path, "rb") as f:
            st.download_button(
                label="⬇️ Download Word Document",
                data=f,
                file_name=f"Method_Statement_{activity_title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("❌ Word file validation failed. Please try regenerating.")
